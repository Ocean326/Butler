from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

out_path = Path("docs/daily-upgrade/0307/20260307_本轮改动与排查总结.docx")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)

doc.add_heading("飞书 管家bot 本轮改动与排查总结", level=1)
doc.add_paragraph("日期：2026-03-07")
doc.add_paragraph("范围：scripts/feishu-bots 本轮实现、排查与验证结果")

doc.add_heading("一、小功能改动", level=2)
items = [
    "运行时模型切换：支持在对话中指定本轮模型，如“用 gpt-5 回答：...”或“[模型=sonnet-4] ...”。",
    "模型查询能力：支持“模型列表”“当前模型”两类查询指令，不需要改配置文件即可查看或指定。",
    "模型别名机制：增加 model_aliases（如 fast/quick -> auto），统一解析后传给 Cursor CLI。",
    "模型列表真实探测：通过 cursor-agent.cmd models --trust --workspace 实时获取可用模型，并解析“model-id - label”格式。",
    "本机调用记忆接口：新增 bots/memory_cli.py，支持 recent-list/recent-add/local-query/local-add/paths。",
    "PowerShell 单次调用支持模型参数：feishu/agent-to-feishu.ps1 增加 -Model 参数。",
    "配置与文档同步：butler_bot.json(.example) 与 README 已更新模型、心跳、记忆 CLI 的使用说明。",
]
for i in items:
    doc.add_paragraph(i, style="List Bullet")

doc.add_heading("二、心跳机制", level=2)
items = [
    "实现位置：bots/memory_manager.py，作为启动后台服务的一部分启动，不引入 OS 独立守护进程。",
    "实现方式：单进程“单次计时器 + 重挂”循环（符合 setTimeout 到期执行后再 arm 的思路）。",
    "调度策略：首次等待 startup_delay_seconds，后续按 every_minutes 周期触发。",
    "发送路径：复用飞书私聊发送能力，支持 heartbeat.receive_id / receive_id_type，未配置时可回退启动通知目标。",
    "可扩展能力：支持 heartbeat.agent_prompt 先调用模型生成心跳内容，再发送。",
]
for i in items:
    doc.add_paragraph(i, style="List Bullet")

doc.add_paragraph("已排查/解决问题：", style=None)
issues = [
    "问题：心跳首次延迟和后续周期可能被混用。",
    "定位：无状态计算会导致每轮都使用 startup_delay_seconds。",
    "解决：引入 heartbeat bootstrap 状态，首次走 startup_delay_seconds，后续固定走 every_minutes。",
    "验证：单测通过（test_heartbeat_uses_startup_delay_once_then_interval）；启动日志出现“[心跳服务] 已启动（单次计时器重挂模式）”。",
]
for i in issues:
    doc.add_paragraph(i, style="List Bullet")

doc.add_heading("三、记忆机制", level=2)
items = [
    "recent_memory 写入可靠性增强：on_reply_sent_async 先同步写 fallback completed 记录，再异步精炼。",
    "避免漏写窗口：即使异步线程异常或进程时序不利，recent_memory.json 也已有可用结果。",
    "记忆接口封装：MemoryManager 增加 get_recent_entries / append_recent_entry / append_local_memory_entry / query_local_memory。",
    "私聊发送通用化：抽出 _send_private_message，启动通知和心跳复用。",
]
for i in items:
    doc.add_paragraph(i, style="List Bullet")

doc.add_paragraph("已排查/解决问题：", style=None)
issues = [
    "问题：偶发 recent_memory 未记录（怀疑子 agent/线程时序导致）。",
    "定位：原流程依赖异步线程完成最终写回，存在落盘竞态窗口。",
    "解决：增加同步 fallback 落盘，再异步 refine 覆盖，且保持 long-term upsert 逻辑不变。",
    "验证：单测通过（test_on_reply_sent_async_writes_fallback_before_refine_finishes），日志可见 [recent-fallback]。",
]
for i in issues:
    doc.add_paragraph(i, style="List Bullet")

doc.add_heading("四、对话机制", level=2)
items = [
    "在 run_agent 入口新增运行时控制解析：区分 list-models / current-model / run。",
    "本轮模型覆盖贯通：解析出的 model 传到 agent 执行与记忆写入链路。",
    "空问题保护：只给了模型指令但没给正文时，返回明确提示示例。",
    "当前模型回复优化：展示默认模型与有效别名。",
]
for i in items:
    doc.add_paragraph(i, style="List Bullet")

doc.add_paragraph("已排查/解决问题：", style=None)
issues = [
    "问题：首次冒烟中“模型列表”返回 0 个模型。",
    "定位：解析器未覆盖 Cursor CLI 的“model-id - label”输出格式。",
    "解决：新增 dash 格式提取逻辑，提取“ - ”前的模型 id。",
    "验证：本地冒烟“模型列表”返回 5 个模型；新增单测 test_list_available_models_parses_dash_format 通过。",
]
for i in issues:
    doc.add_paragraph(i, style="List Bullet")

doc.add_heading("五、测试与运行状态", level=2)
items = [
    "回归测试：tests.test_butler_bot_model_controls / tests.test_memory_manager_recent / tests.test_memory_manager_maintenance / tests.test_butler_bot_streaming 全部通过。",
    "总计：Ran 19 tests, OK。",
    "本机记忆 CLI 冒烟通过：python bots/memory_cli.py paths --json 返回路径正确。",
    "服务状态：已重启管家bot（butler_bot），当前进程正常运行。",
]
for i in items:
    doc.add_paragraph(i, style="List Bullet")

doc.add_heading("六、涉及文件（本轮核心）", level=2)
files = [
    "bots/butler_bot.py",
    "bots/memory_manager.py",
    "bots/memory_cli.py",
    "feishu/agent-to-feishu.ps1",
    "configs/butler_bot.json",
    "configs/butler_bot.json.example",
    "README.md",
    "tests/test_butler_bot_model_controls.py",
    "tests/test_butler_bot_streaming.py",
    "tests/test_memory_manager_recent.py",
    "tests/test_memory_manager_maintenance.py",
]
for f in files:
    doc.add_paragraph(f, style="List Bullet")

doc.add_paragraph("生成时间：" + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
doc.save(out_path)
print(out_path)
